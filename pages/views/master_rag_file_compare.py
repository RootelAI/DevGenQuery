import os
import traceback
from django.shortcuts import render
from azure.storage.blob import BlobServiceClient
from docx import Document
from sentence_transformers import SentenceTransformer, util
from utilsPrj.supabase_client import get_supabase_client
import re
from io import BytesIO

# 서버 시작 시 1회만 모델 로딩
model = SentenceTransformer('all-MiniLM-L6-v2')

# ==============================
# 섹션 추출 (Word)
# ==============================
def extract_sections_from_docx(blob_bytes):
    doc = Document(BytesIO(blob_bytes))
    sections = {}
    current_title = "PREAMBLE"
    sections[current_title] = {"text": "", "start_para": 0}

    para_id = 0
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            current_title = para.text.strip()
            sections[current_title] = {"text": "", "start_para": para_id}
        else:
            sections[current_title]["text"] += para.text + " "
        para_id += 1

    for sec in sections.values():
        sec["text"] = sec["text"].strip()
    return sections

# ==============================
# 문장 분리
# ==============================
def split_into_sentences(text):
    text = re.sub(r'\s+', ' ', text).strip()
    sentences = re.split(r'(?<=[.?!])\s+', text)
    return [s.strip() for s in sentences if s.strip()]

# ==============================
# 문장 단위 비교
# ==============================
def compare_sentences(left_sents, right_sents, threshold=0.80):
    rows = []
    if not left_sents or not right_sents:
        return rows

    emb_left = model.encode(left_sents, convert_to_tensor=True)
    emb_right = model.encode(right_sents, convert_to_tensor=True)
    cos_scores = util.cos_sim(emb_left, emb_right)

    # 왼쪽 기준 비교
    for i, sent in enumerate(left_sents):
        max_idx = cos_scores[i].argmax().item()
        max_sim = cos_scores[i, max_idx].item()
        if max_sim < threshold:
            rows.append({
                "side": "과거",
                "sentence": sent,
                "compare_sentence": right_sents[max_idx],
                "similarity": f"{max_sim:.3f}"
            })

    # 오른쪽에서 매칭 안 된 문장
    for j, sent in enumerate(right_sents):
        if cos_scores[:, j].max().item() < threshold:
            rows.append({
                "side": "현재",
                "sentence": sent,
                "compare_sentence": "",
                "similarity": "-"
            })
    return rows

# ==============================
# 섹션 단위 매칭 + 수정 섹션만 문장 비교
# ==============================
def match_sections_vector_db(sections1, sections2, threshold=0.80):
    diff_list = []

    titles1 = list(sections1.keys())
    titles2 = list(sections2.keys())

    # 제목+본문 전체를 합쳐 벡터 생성
    emb_titles1 = model.encode([t + " " + sections1[t]["text"] for t in titles1], convert_to_tensor=True)
    emb_titles2 = model.encode([t + " " + sections2[t]["text"] for t in titles2], convert_to_tensor=True)

    cos_scores = util.cos_sim(emb_titles1, emb_titles2)
    matched_2 = set()

    for i, title1 in enumerate(titles1):
        max_idx = cos_scores[i].argmax().item()
        max_sim = cos_scores[i, max_idx].item()
        title2 = titles2[max_idx]

        if max_sim >= threshold:
            # 수정 섹션만 문장 단위 비교
            left_sents = split_into_sentences(sections1[title1]["text"])
            right_sents = split_into_sentences(sections2[title2]["text"])
            rows = compare_sentences(left_sents, right_sents)
            change_type = "동일" if not rows else "내용 변경"

            diff_list.append({
                "section_title_old": title1,
                "section_title_new": title2,
                "change_type": change_type,
                "rows": rows
            })
            matched_2.add(max_idx)
        else:
            # 삭제된 섹션
            diff_list.append({
                "section_title_old": title1,
                "section_title_new": "-",
                "change_type": "삭제",
                "rows": []
            })

    # 새 문서에서 매칭 안 된 섹션 → 신규
    for j, title2 in enumerate(titles2):
        if j not in matched_2:
            diff_list.append({
                "section_title_old": "-",
                "section_title_new": title2,
                "change_type": "신규",
                "rows": []
            })

    return diff_list

# ==============================
# Word 파일 비교 뷰
# ==============================
def master_rag_file_compare(request):
    access_token = request.session.get("access_token")
    refresh_token = request.session.get("refresh_token")
    user = request.session.get("user")

    if not user:
        return render(request, "pages/home.html", {"code": "login", "text": "로그인이 필요합니다."})

    try:
        project_id = 3
        supabase = get_supabase_client(access_token, refresh_token)
        dirPath = supabase.schema('rag').table('projects') \
            .select('dirpath').eq('projectid', project_id) \
            .execute().data[0]['dirpath']

        blob_service_client = BlobServiceClient.from_connection_string(
            os.getenv("AZURE_STORAGE_CONNECTION_STRING")
        )
        container_client = blob_service_client.get_container_client(dirPath)

        blob1 = container_client.download_blob("source/문서_01.docx").readall()
        blob2 = container_client.download_blob("source/문서_02.docx").readall()

        # 섹션 추출
        sections1 = extract_sections_from_docx(blob1)
        sections2 = extract_sections_from_docx(blob2)

        # 섹션 비교
        diff_list = match_sections_vector_db(sections1, sections2)

        # Word -> HTML 변환용
        html_left = "".join(f"<p>{p.text}</p>" for p in Document(BytesIO(blob1)).paragraphs)
        html_right = "".join(f"<p>{p.text}</p>" for p in Document(BytesIO(blob2)).paragraphs)

        return render(request, "pages/master_rag_file_compare.html", {
            "diff_list": diff_list,
            "html_left": html_left,
            "html_right": html_right
        })

    except Exception as e:
        traceback.print_exc()
        return render(request, "pages/master_rag_file_compare.html", {"error": str(e)})