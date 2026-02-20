import os
import io
import tempfile
import fitz  # PyMuPDF
from dotenv import load_dotenv, dotenv_values

from azure.storage.blob import BlobServiceClient

from langchain_openai.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from supabase import create_client
from datetime import datetime, timezone, timedelta
from docx import Document as DocxDocument

# ==============================
# 1. 환경변수 로드
# ==============================
if os.path.exists(".env"):
    load_dotenv(override=True)
    config = dotenv_values(".env")

											  
AZURE_CONNECTION_STRING = os.getenv("AZURE_STORAGE_CONNECTION_STRING")

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")			
															   
SOURCE_PREFIX = "source/"
VECTORDB_PREFIX = "vectordb/"
embeddings_model = "text-embedding-3-large"

embeddings = OpenAIEmbeddings(model=embeddings_model)

blob_service_client = BlobServiceClient.from_connection_string(
    AZURE_CONNECTION_STRING
)

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

# ==========================================================							
# 핵심 함수: dirpath(=container_name) 만 재적재
# ==========================================================
def rebuild_vectordb(dirpath: str):

    if not dirpath:
        raise ValueError("dirpath가 필요합니다.")

    # if dirpath not in ["law", "medicine"]:
    #     raise ValueError(f"지원하지 않는 dirpath 입니다: {dirpath}")

    container_name = dirpath

    # print(f"\n==============================")
    # print(f"컨테이너 처리 시작: {container_name}")
    # print(f"==============================")

    container_client = blob_service_client.get_container_client(container_name)

    documents = []

    # --------------------------------
    # 1. Azure source/ 내 PDF 조회
    # --------------------------------
    blob_list = container_client.list_blobs(name_starts_with=SOURCE_PREFIX)

    for blob in blob_list:
        filename = os.path.basename(blob.name)
        blob_client = container_client.get_blob_client(blob.name)
        blob_data = blob_client.download_blob().readall()
        # print(f"Processing Blob: {blob.name}")

        # --------------------------------
        # 2. Supabase 메타데이터 조회
        # --------------------------------
        file_response = (
					
            supabase.schema('rag')
            .table('files')
            .select('filestatus, filemastercd')
            .eq('filenm', filename)
            .execute()
        )

        file_data = file_response.data[0] if file_response.data else {}
        filestatus = file_data.get("filestatus")
        filemastercd = file_data.get("filemastercd")

        filemaster_data = {}
        projectid = None
        owner_dept = None

        if filemastercd:
            filemaster_response = (
						
                supabase.schema('rag')
                .table('filemasters')
                .select('projectid, owner_dept, tag1value, tag2value, tag3value, tag4value, tag5value')
                .eq('filemastercd', filemastercd)
                .execute()
            )

            if filemaster_response.data:
                filemaster_data = filemaster_response.data[0]
                projectid = filemaster_data.get("projectid")
                owner_dept = filemaster_data.get("owner_dept")

        projecttags = []
        if projectid:
            projecttag_response = (
						
                supabase.schema('rag')
                .table('projecttags')
                .select('tagcd, tagnm')
                .eq('projectid', projectid)
                .execute()
            )
            projecttags = projecttag_response.data or []

        metadata_tags = {
            "주관부서": owner_dept,
            "파일상태": filestatus,
            "container": container_name
        }

        for i, tag in enumerate(projecttags, start=1):
            tagnm = tag["tagnm"]
            tag_value = filemaster_data.get(f"tag{i}value")
            if tag_value is not None:
                metadata_tags[tagnm] = tag_value

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100
        )

        # --------------------------------
        # 3. PDF 처리
        # --------------------------------
        if blob.name.endswith(".pdf"):
            pdf_stream = io.BytesIO(blob_data)

            try:
                doc = fitz.open(stream=pdf_stream, filetype="pdf")
            except Exception as e:
                # print(f"PDF 열기 실패: {e}")
                continue

            for page_number, page in enumerate(doc):
                page_text = page.get_text()

                chunks = text_splitter.create_documents([page_text])

                for chunk in chunks:
                    documents.append(
                        Document(
                            page_content=chunk.page_content,
                            metadata={
                                "source": blob.name,
                                "page": page_number,
                                **metadata_tags
                            }
                        )
                    )

            doc.close()

        # --------------------------------
        # 4. DOCX 처리
        # --------------------------------
        elif blob.name.endswith(".docx"):
            try:
                docx_stream = io.BytesIO(blob_data)
                docx_doc = DocxDocument(docx_stream)
                full_text = "\n".join([p.text for p in docx_doc.paragraphs])
            except Exception:
                continue

            chunks = text_splitter.create_documents([full_text])
            for i, chunk in enumerate(chunks):
                documents.append(
                    Document(
                        page_content=chunk.page_content,
                        metadata={"source": blob.name, "page": i, **metadata_tags}
                    )
                )

        else:
            # PDF, DOCX 외 파일은 건너뜀
            continue

    # print(f"{container_name} → 총 문서 청크 수: {len(documents)}")

    if not documents:
        # print("문서 없음.")
        return

    # --------------------------------
    # 5. FAISS 생성
    # --------------------------------
    # print("FAISS DB 생성 중...")

    db = FAISS.from_documents(
        documents=documents,
        embedding=embeddings,
        normalize_L2=True
    )

    # --------------------------------
    # 6. 기존 vectordb 삭제
    # --------------------------------
    # print("기존 vectordb 삭제 중...")

    existing_blobs = container_client.list_blobs(name_starts_with=VECTORDB_PREFIX)

    for blob in existing_blobs:
        container_client.delete_blob(blob.name)

    # print("기존 vectordb 삭제 완료")

    # --------------------------------
    # 7. Azure vectordb 업로드
    # --------------------------------
    with tempfile.TemporaryDirectory() as tmpdir:

        local_path = os.path.join(tmpdir, "faiss_db")
        db.save_local(local_path)

        for filename in os.listdir(local_path):
            file_path = os.path.join(local_path, filename)

            blob_name = f"{VECTORDB_PREFIX}{filename}"
            blob_client = container_client.get_blob_client(blob_name)

            with open(file_path, "rb") as data:
                blob_client.upload_blob(data, overwrite=True)

    # print(f"{container_name} → vectordb 업로드 완료!")

    # --------------------------------
    # 8. processcd 업데이트
    # --------------------------------
    kst = timezone(timedelta(hours=9))
    processcd_data = {
        "processcd": "Y",
        # "processdts": datetime.now(timezone.utc).isoformat()
        "processdts": datetime.now(kst).isoformat()
    }

    # 1️⃣ 해당 project의 filemastercd 조회
    filemaster_resp = (
        supabase.schema("rag")
        .table("filemasters")
        .select("filemastercd")
        .eq("processcd", "N")
        .eq("projectid", projectid)
        .execute()
    )

    filemaster_list = filemaster_resp.data or []

    if not filemaster_list:
        pass
        # print("업데이트할 filemaster 없음")
    else:
        filemastercd_list = [fm["filemastercd"] for fm in filemaster_list]

        # 2️⃣ filemasters 업데이트
        supabase.schema("rag") \
            .table("filemasters") \
            .update(processcd_data) \
            .in_("filemastercd", filemastercd_list) \
            .execute()

        # 3️⃣ files 업데이트 (해당 filemaster만)
        supabase.schema("rag") \
            .table("files") \
            .update(processcd_data) \
            .in_("filemastercd", filemastercd_list) \
            .execute()

    # print(f"{container_name} → DB 업데이트 완료")
