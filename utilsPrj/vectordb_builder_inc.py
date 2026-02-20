# utilsPrj/vectordb_builder_inc.py
import os
import io
import tempfile
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from dotenv import load_dotenv, dotenv_values

from azure.storage.blob import BlobServiceClient

from langchain_openai.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from supabase import create_client
from datetime import datetime, timezone, timedelta

# ==============================
# 1. 환경변수 로드
# ==============================
if os.path.exists(".env"):
    load_dotenv(override=True)
    config = dotenv_values(".env")

AZURE_CONNECTION_STRING = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")			

SOURCE_PREFIX = "source/"
VECTORDB_PREFIX = "vectordb/"
embeddings_model = "text-embedding-3-large"

embeddings = OpenAIEmbeddings(model=embeddings_model)
blob_service_client = BlobServiceClient.from_connection_string(AZURE_CONNECTION_STRING)
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

# utilsPrj/vectordb_builder_inc.py

def rebuild_vectordb_incremental(dirpath: str):
    if not dirpath:
        raise ValueError("dirpath가 필요합니다.")

    container_name = dirpath
    container_client = blob_service_client.get_container_client(container_name)
    documents = []

    # -------------------------
    # 1️⃣ projects 조회 (dirpath 기준)
    # -------------------------
    project_resp = (
        supabase.schema("rag")
        .table("projects")
        .select("projectid")
        .eq("dirpath", container_name)
        .execute()
    )

    projects = project_resp.data or []
    if not projects:
        return

    project_ids = [p["projectid"] for p in projects]

    # -------------------------
    # 2️⃣ filemasters 조회 (projectid 기준)
    # -------------------------
    filemasters_resp = (
        supabase.schema("rag")
        .table("filemasters")
        .select(
            "filemastercd, projectid, owner_dept, "
            "tag1value, tag2value, tag3value, tag4value, tag5value, processcd"
        )
        .in_("projectid", project_ids)
        .execute()
    )

    filemasters = filemasters_resp.data or []
    if not filemasters:
        return

    filemastercds = [fm["filemastercd"] for fm in filemasters]
    filemasters_map = {fm["filemastercd"]: fm for fm in filemasters}

    # -------------------------
    # 3️⃣ files 조회 (filemastercd 기준)
    # -------------------------
    files_resp = (
        supabase.schema("rag")
        .table("files")
        .select("filecd, filenm, filemastercd, filestatus, processcd")
        .in_("filemastercd", filemastercds)
        .execute()
    )

    all_files = files_resp.data or []
    if not all_files:
        return

    # -------------------------
    # 4️⃣ 처리 대상 필터링
    # -------------------------
    files_to_process = []

    for f in all_files:
        fm = filemasters_map.get(f.get("filemastercd"))
        fm_processcd = fm.get("processcd") if fm else "Y"
        file_processcd = f.get("processcd", "Y")

        if fm_processcd != "Y" or file_processcd != "Y":
            files_to_process.append(f)

    if not files_to_process:
        return

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100
    )
    
    # -------------------------
    # 5️⃣ PDF 읽기 + Chunk 생성
    # -------------------------
    for file_item in files_to_process:
        filename = file_item["filenm"]
        filestatus = file_item.get("filestatus")
        filemastercd = file_item.get("filemastercd")

        fm_data = filemasters_map.get(filemastercd, {})
        projectid = fm_data.get("projectid")
        owner_dept = fm_data.get("owner_dept")

        # 프로젝트 태그 조회
        projecttags = []
        if projectid:
            pt_resp = (
                supabase.schema("rag")
                .table("projecttags")
                .select("tagcd, tagnm")
                .eq("projectid", projectid)
                .execute()
            )
            projecttags = pt_resp.data or []

        metadata_tags = {
            "주관부서": owner_dept,
            "파일상태": filestatus,
            "container": container_name,
        }

        for i, tag in enumerate(projecttags, start=1):
            tag_value = fm_data.get(f"tag{i}value")
            if tag_value is not None:
                metadata_tags[tag["tagnm"]] = tag_value

        blob_client = container_client.get_blob_client(f"{SOURCE_PREFIX}{filename}")
        blob_data = blob_client.download_blob().readall()

        if filename.lower().endswith(".pdf"):
            try:
                pdf_stream = io.BytesIO(blob_data)
                doc = fitz.open(stream=pdf_stream, filetype="pdf")
            except Exception:
                continue

            for page_number, page in enumerate(doc):
                chunks = text_splitter.create_documents([page.get_text()])
                for chunk in chunks:
                    documents.append(
                        Document(
                            page_content=chunk.page_content,
                            metadata={
                                "source": filename,
                                "page": page_number,
                                **metadata_tags,
                            },
                        )
                    )
            doc.close()
        # ---------------- DOCX ----------------
        elif filename.lower().endswith(".docx"):
            try:
                docx_stream = io.BytesIO(blob_data)
                docx_doc = DocxDocument(docx_stream)
                full_text = "\n".join([p.text for p in docx_doc.paragraphs])
            except Exception:
                continue

            chunks = text_splitter.create_documents([full_text])
            for idx, chunk in enumerate(chunks):
                documents.append(
                    Document(
                        page_content=chunk.page_content,
                        metadata={
                            "source": filename,
                            "page": idx,
                            **metadata_tags,
                        },
                    )
                )

        else:
            continue

    if not documents:
        return

    # -------------------------
    # 6️⃣ FAISS 생성/병합
    # -------------------------
    local_db_path = f"/tmp/{container_name}_faiss"

    try:
        db = FAISS.load_local(local_db_path, embeddings)
        db.add_documents(documents)
    except Exception:
        db = FAISS.from_documents(documents, embeddings, normalize_L2=True)

    # -------------------------
    # 7️⃣ Azure 업로드
    # -------------------------
    with tempfile.TemporaryDirectory() as tmpdir:
        db_path = os.path.join(tmpdir, "faiss_db")
        db.save_local(db_path)

        for fname in os.listdir(db_path):
            fpath = os.path.join(db_path, fname)
            blob_client = container_client.get_blob_client(
                f"{VECTORDB_PREFIX}{fname}"
            )
            with open(fpath, "rb") as f:
                blob_client.upload_blob(f, overwrite=True)

    # -------------------------
    # 8️⃣ processcd 업데이트
    # -------------------------
    kst = timezone(timedelta(hours=9))
    processcd_data = {
        "processcd": "Y",
        "processdts": datetime.now(kst).isoformat(),
    }

    # filemasters 업데이트
    filemastercd_list = list({
        f["filemastercd"]
        for f in files_to_process
        if f.get("filemastercd")
    })

    if filemastercd_list:
        supabase.schema("rag").table("filemasters") \
            .update(processcd_data) \
            .in_("filemastercd", filemastercd_list) \
            .neq("processcd", "Y") \
            .execute()

    # files 업데이트 (filecd 기준 안전)
    filecd_list = list({
        f["filecd"]
        for f in files_to_process
        if f.get("filecd")
    })

    if filecd_list:
        supabase.schema("rag").table("files") \
            .update(processcd_data) \
            .in_("filecd", filecd_list) \
            .neq("processcd", "Y") \
            .execute()
